import pytest
from docx import Document

import mcp_server_qa as server


# --- adf_to_text ---


def test_adf_to_text_plain_paragraph():
    node = {"type": "paragraph", "content": [{"type": "text", "text": "Hola mundo"}]}
    assert server.adf_to_text(node) == "Hola mundo\n"


def test_adf_to_text_heading_and_list():
    node = {
        "type": "doc",
        "content": [
            {"type": "heading", "content": [{"type": "text", "text": "Título"}]},
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [{"type": "text", "text": "Item 1"}],
                    }
                ],
            },
        ],
    }
    assert server.adf_to_text(node) == "Título\n- Item 1\n"


def test_adf_to_text_none_returns_empty():
    assert server.adf_to_text(None) == ""


# --- get_target_content ---


def _make_docx(tmp_path, parrafos):
    """parrafos: lista de párrafos, cada uno una lista de (texto, bold)."""
    doc = Document()
    for runs in parrafos:
        p = doc.add_paragraph()
        for texto, bold in runs:
            run = p.add_run(texto)
            run.bold = bold
    path = tmp_path / "target.docx"
    doc.save(path)
    return str(path)


def test_get_target_content_reports_bold_state(tmp_path):
    doc_path = _make_docx(
        tmp_path, [[("Hola ", False), ("mundo", True)], [("Otro párrafo", False)]]
    )
    resultado = server.get_target_content(doc_path)
    assert '"Hola " (sin negrita)' in resultado
    assert '"mundo" (negrita)' in resultado
    assert "Párrafo 2:" in resultado


def test_get_target_content_missing_file():
    resultado = server.get_target_content("/no/existe/archivo.docx")
    assert "No se pudo leer el documento" in resultado


def test_get_target_content_empty_doc(tmp_path):
    doc = Document()
    path = tmp_path / "vacio.docx"
    doc.save(path)
    resultado = server.get_target_content(str(path))
    assert resultado == "El documento está vacío o no tiene texto legible."


# --- get_ticket ---


class FakeResponse:
    def __init__(self, status_code, json_data=None, text=""):
        self.status_code = status_code
        self._json_data = json_data or {}
        self.text = text

    def json(self):
        return self._json_data


def test_get_ticket_missing_credentials(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", None)
    monkeypatch.setattr(server, "JIRA_EMAIL", None)
    monkeypatch.setattr(server, "JIRA_API_TOKEN", None)
    resultado = server.get_ticket("SCRUM-1")
    assert "Faltan credenciales" in resultado


def test_get_ticket_not_found(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", "example.atlassian.net")
    monkeypatch.setattr(server, "JIRA_EMAIL", "a@b.com")
    monkeypatch.setattr(server, "JIRA_API_TOKEN", "token")
    monkeypatch.setattr(server.requests, "get", lambda *a, **k: FakeResponse(404))
    resultado = server.get_ticket("SCRUM-999")
    assert "No se encontró ningún ticket" in resultado


def test_get_ticket_unauthorized(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", "example.atlassian.net")
    monkeypatch.setattr(server, "JIRA_EMAIL", "a@b.com")
    monkeypatch.setattr(server, "JIRA_API_TOKEN", "token")
    monkeypatch.setattr(server.requests, "get", lambda *a, **k: FakeResponse(401))
    resultado = server.get_ticket("SCRUM-1")
    assert "No autorizado" in resultado


def test_get_ticket_success(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", "example.atlassian.net")
    monkeypatch.setattr(server, "JIRA_EMAIL", "a@b.com")
    monkeypatch.setattr(server, "JIRA_API_TOKEN", "token")
    fields = {
        "summary": "Arreglar botón",
        "description": {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Debe ser azul"}],
                }
            ],
        },
    }
    monkeypatch.setattr(
        server.requests,
        "get",
        lambda *a, **k: FakeResponse(200, json_data={"fields": fields}),
    )
    resultado = server.get_ticket("SCRUM-1")
    assert "Título: Arreglar botón" in resultado
    assert "Debe ser azul" in resultado


# --- add_ticket_comment ---


def test_text_to_adf_one_paragraph_per_nonblank_line():
    adf = server._text_to_adf("Cumple.\n\nRazón: el botón es azul.")
    assert adf == {
        "type": "doc",
        "version": 1,
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "Cumple."}]},
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Razón: el botón es azul."}
                ],
            },
        ],
    }


def test_text_to_adf_empty_text_returns_empty_paragraph():
    assert server._text_to_adf("") == {
        "type": "doc",
        "version": 1,
        "content": [{"type": "paragraph", "content": []}],
    }


def test_add_ticket_comment_missing_credentials(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", None)
    monkeypatch.setattr(server, "JIRA_EMAIL", None)
    monkeypatch.setattr(server, "JIRA_API_TOKEN", None)
    resultado = server.add_ticket_comment("SCRUM-1", "Cumple.")
    assert "Faltan credenciales" in resultado


def test_add_ticket_comment_not_found(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", "example.atlassian.net")
    monkeypatch.setattr(server, "JIRA_EMAIL", "a@b.com")
    monkeypatch.setattr(server, "JIRA_API_TOKEN", "token")
    monkeypatch.setattr(server.requests, "post", lambda *a, **k: FakeResponse(404))
    resultado = server.add_ticket_comment("SCRUM-999", "Cumple.")
    assert "No se encontró ningún ticket" in resultado


def test_add_ticket_comment_forbidden(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", "example.atlassian.net")
    monkeypatch.setattr(server, "JIRA_EMAIL", "a@b.com")
    monkeypatch.setattr(server, "JIRA_API_TOKEN", "token")
    monkeypatch.setattr(server.requests, "post", lambda *a, **k: FakeResponse(403))
    resultado = server.add_ticket_comment("SCRUM-1", "Cumple.")
    assert "permisos" in resultado


def test_add_ticket_comment_success(monkeypatch):
    monkeypatch.setattr(server, "JIRA_DOMAIN", "example.atlassian.net")
    monkeypatch.setattr(server, "JIRA_EMAIL", "a@b.com")
    monkeypatch.setattr(server, "JIRA_API_TOKEN", "token")

    captured = {}

    def fake_post(url, auth=None, headers=None, json=None, timeout=None):
        captured["json"] = json
        return FakeResponse(201, json_data={"id": "10042"})

    monkeypatch.setattr(server.requests, "post", fake_post)
    resultado = server.add_ticket_comment("SCRUM-1", "Cumple.")
    assert "Comentario agregado al ticket SCRUM-1" in resultado
    assert "10042" in resultado
    assert captured["json"]["body"]["content"][0]["content"][0]["text"] == "Cumple."


# --- _open_page_with_element (helper compartido por las tools de Playwright) ---


class FakeLocator:
    def __init__(self, visible=True):
        self.visible = visible

    @property
    def first(self):
        return self

    async def wait_for(self, state="visible", timeout=0):
        if not self.visible:
            raise TimeoutError("elemento no visible")


class FakePage:
    def __init__(self, locator):
        self._locator = locator
        self.closed = False
        self.url = "about:blank"

    def locator(self, selector):
        return self._locator

    async def goto(self, url, wait_until="load", timeout=0):
        self.url = url

    async def close(self):
        self.closed = True


class FakeBrowser:
    def __init__(self, page):
        self._page = page

    async def new_page(self):
        return self._page


async def test_open_page_with_element_yields_page_and_locator():
    locator = FakeLocator(visible=True)
    page = FakePage(locator)
    browser = FakeBrowser(page)

    async with server._open_page_with_element(browser, "http://x", "#sel") as (
        p,
        loc,
    ):
        assert p is page
        assert loc is locator
    assert page.closed is True


async def test_open_page_with_element_raises_and_closes_page_when_not_found():
    locator = FakeLocator(visible=False)
    page = FakePage(locator)
    browser = FakeBrowser(page)

    with pytest.raises(server.ElementNotFoundError):
        async with server._open_page_with_element(browser, "http://x", "#sel"):
            pass
    assert page.closed is True
